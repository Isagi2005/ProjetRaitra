# utils.py

from docx import Document
from docx.shared import Pt
from django.http import HttpResponse
from datetime import date
import io

def creer_document_certificat(nom_etudiant: str, texte_certificat: str) -> HttpResponse:
    doc = Document()

    titre = doc.add_heading("RAITRA KIDZ", level=0)
    titre.alignment = 1

    sous_titre = doc.add_paragraph("By Pass Alasora")
    sous_titre.alignment = 1

    doc.add_paragraph().add_run().add_break()

    cert_title = doc.add_paragraph()
    cert_title.alignment = 1
    run = cert_title.add_run("CERTIFICAT DE SCOLARITÉ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph().add_run().add_break()

    corps = doc.add_paragraph(texte_certificat)
    corps.alignment = 3
    corps.style.font.size = Pt(12)

    doc.add_paragraph().add_run().add_break()

    today = date.today().strftime("%d/%m/%Y")
    doc.add_paragraph(f"Fait à Alasora, le {today}")

    doc.add_paragraph().add_run().add_break()
    signature = doc.add_paragraph("Signature et cachet de la direction")
    signature.alignment = 2

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename=certificat_{nom_etudiant}.docx'

    return response
