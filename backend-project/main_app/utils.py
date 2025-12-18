import logging
import traceback
from django.core.mail import send_mail
from django.conf import settings
import google.generativeai as genai
import os
from dotenv import load_dotenv

# Initialisation du logger
logger = logging.getLogger(__name__)

# Charger les variables d'environnement
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")

# Configuration de Gemini
try:
    genai.configure(api_key=API_KEY)
    logger.info("Gemini API configurée avec succès.")
except Exception as e:
    logger.error(f"Erreur lors de la configuration de Gemini : {e}")
    raise

# --- Partie EMAIL ---
def envoyer_email_notification(email_to, subject, message):
    try:
        if not all([settings.EMAIL_HOST, settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD]):
            logger.error("Configuration email incomplète dans settings.py")
            return None

        logger.info(f"Tentative d'envoi d'email à {email_to}")
        logger.debug(f"Paramètres: HOST={settings.EMAIL_HOST}, USER={settings.EMAIL_HOST_USER}")

        email_id = send_mail(
            subject=subject,
            message=message,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[email_to],
            fail_silently=False,
        )

        logger.info(f"Email envoyé avec succès à {email_to}, ID: {email_id}")
        return email_id

    except Exception as e:
        logger.error(f"Erreur lors de l'envoi de l'email : {str(e)}")
        logger.error(traceback.format_exc())
        return None

import logging
import traceback
import google.generativeai as genai
import os
from dotenv import load_dotenv
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from datetime import date
import io
# Configuration initiale
load_dotenv()
logger = logging.getLogger(__name__)

def configurer_gemini():
    try:
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        return genai.GenerativeModel("gemini-1.5-pro-latest")
    except Exception as e:
        logger.error(f"Erreur configuration Gemini: {str(e)}")
        raise

def generer_texte_certificat(prompt: str) -> str:
    try:
        model = configurer_gemini()
        response = model.generate_content(prompt)
        
        if not response.text:
            raise ValueError("Réponse vide de Gemini")
            
        return response.text
        
    except Exception as e:
        logger.error(f"Erreur génération texte: {traceback.format_exc()}")
        raise
def creer_document_certificat(nom_etudiant: str, texte_certificat: str) -> HttpResponse:
    """Crée un document Word pour le certificat de scolarité"""
    try:
        doc = Document()

        # En-tête
        titre = doc.add_heading("RAITRA KIDZ", level=0)
        titre.alignment = 1  # Centré

        sous_titre = doc.add_paragraph("By Pass Alasora")
        sous_titre.alignment = 1  # Centré

        doc.add_paragraph().add_run().add_break()  # Saut de ligne

        # Titre du certificat
        cert_title = doc.add_paragraph()
        cert_title.alignment = 1
        run = cert_title.add_run("CERTIFICAT DE SCOLARITÉ")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph().add_run().add_break()

        # Corps du texte
        corps = doc.add_paragraph(texte_certificat)
        corps.alignment = 3  # Justifié
        corps.style.font.size = Pt(12)

        doc.add_paragraph().add_run().add_break()

        # Pied de page
        today = date.today().strftime("%d/%m/%Y")
        doc.add_paragraph(f"Fait à Alasora, le {today}")

        doc.add_paragraph().add_run().add_break()
        signature = doc.add_paragraph("Signature et cachet de la direction")
        signature.alignment = 2  # Alignement à droite

        # Préparation de la réponse HTTP
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename=certificat_{nom_etudiant}.docx'

        return response

    except Exception as e:
        logger.error(f"Erreur lors de la création du document: {str(e)}")
        logger.error(traceback.format_exc())
        raise